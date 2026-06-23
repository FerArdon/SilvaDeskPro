"""
utils/word_contrato.py – Generador de Contrato Word (.docx)
SilvaDesk Pro · SEDCAF / ICF Honduras
Espeja la misma estructura legal de pdf_contrato.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Paleta ────────────────────────────────────────────────────────────────
_VERDE_OSC = RGBColor(0x1B, 0x5E, 0x20)
_VERDE_MED = RGBColor(0x2E, 0x7D, 0x32)
_GRIS      = RGBColor(0x33, 0x33, 0x33)
_NEGRO     = RGBColor(0x00, 0x00, 0x00)


# ─── Helpers ────────────────────────────────────────────────────────────────
def _set_run_font(run, bold=False, italic=False, size=10,
                  color=None, name="Times New Roman"):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = color


def _para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
          bold=False, size=10, color=None,
          space_before=0, space_after=4, name="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(14)
    if text:
        run = p.add_run(text)
        _set_run_font(run, bold=bold, size=size, color=color, name=name)
    return p


def _para_mixed(doc, fragments, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6, size=10):
    """
    fragments: list of (text, bold) tuples
    """
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(15)
    for text, bold in fragments:
        run = p.add_run(text)
        _set_run_font(run, bold=bold, size=size, color=_GRIS)
    return p


def _clause_heading(doc, text):
    p = _para(doc, text, bold=True, size=10, color=_NEGRO,
              space_before=8, space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    return p


def _add_hr(doc):
    """Línea horizontal fina entre secciones."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def _footer(doc):
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(
            "SilvaDesk Pro — Contrato de Servicios Profesionales Forestales"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = "Calibri"


# ─── Función principal ───────────────────────────────────────────────────────
def generar_word_contrato(ruta_docx: str, datos: dict) -> str:
    os.makedirs(
        os.path.dirname(ruta_docx) if os.path.dirname(ruta_docx) else ".",
        exist_ok=True
    )

    # Extraer variables (mismas keys que pdf_contrato.py)
    c_nom    = datos.get("consultor_nombre",   "Ing. Fernando Rafael Ardon Rodriguez")
    c_reg    = datos.get("consultor_registro", "COLPROFORH N.- 0226")
    cl_nom   = datos.get("cliente_nombre",     "")
    cl_rtn   = datos.get("cliente_rtn",        "")
    predio   = datos.get("predio",             "")
    municipio= datos.get("municipio",          "")
    depto    = datos.get("departamento",       "")
    actividad= datos.get("actividad_tipo",     "")
    monto    = datos.get("monto_total",        0.0)
    forma_p  = datos.get("forma_pago",         "")
    plazo    = datos.get("plazo_ejecucion",    "")
    fecha    = datos.get("fecha_firma",        "")
    clausulas_extra = datos.get("clausulas_adicionales", "")

    monto_str = f"L. {float(monto):,.2f}"

    doc = Document()
    _set_margins(doc)
    _footer(doc)

    # ── Título principal ────────────────────────────────────────────────────
    t = _para(
        doc,
        "CONTRATO PRIVADO DE PRESTACIÓN DE SERVICIOS PROFESIONALES FORESTALES",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=_NEGRO,
        space_before=0, space_after=2, name="Times New Roman"
    )

    _para(
        doc,
        "CONFORME AL ARANCEL DE PROFESIONALES DE LAS CIENCIAS FORESTALES DE HONDURAS",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, color=_VERDE_MED,
        space_before=0, space_after=10, name="Times New Roman"
    )

    _add_hr(doc)

    # ── Comparecientes ──────────────────────────────────────────────────────
    rtn_display = cl_rtn if cl_rtn else "N/A"
    _para_mixed(doc, [
        ("Nosotros, por una parte, ", False),
        (c_nom, True),
        (", Ingeniero Forestal de nacionalidad hondureña, "
         "miembro activo del Colegio de Profesionales Forestales de Honduras (COLPROFORH) "
         "con registro número ", False),
        (c_reg, True),
        (", en adelante denominado el ", False),
        ("CONSULTOR", True),
        ("; y por otra parte, ", False),
        (cl_nom, True),
        (", mayor de edad, con Registro Tributario Nacional (RTN) número ", False),
        (rtn_display, True),
        (", actuando en su condición de propietario o representante legal de la propiedad, "
         "en adelante denominado el ", False),
        ("PROPIETARIO", True),
        ("; convenimos de mutuo acuerdo celebrar el presente Contrato de Prestación de "
         "Servicios Profesionales Forestales, el cual se regirá por las cláusulas siguientes:", False),
    ], space_after=10)

    # ── CLÁUSULA PRIMERA ────────────────────────────────────────────────────
    _clause_heading(doc, "CLÁUSULA PRIMERA: OBJETO DEL CONTRATO")
    _para_mixed(doc, [
        ("El ", False), ("CONSULTOR", True),
        (" se compromete a realizar para el ", False), ("PROPIETARIO", True),
        (" la asistencia técnica y ejecución del servicio forestal de: ", False),
        (actividad, True),
        (". Dicha labor técnica se desarrollará en el predio denominado ", False),
        (f'"{predio}"', True),
        (", ubicado en el municipio de ", False),
        (municipio if municipio else "______", True),
        (", departamento de ", False),
        (depto if depto else "______", True),
        (", República de Honduras.", False),
    ])

    # ── CLÁUSULA SEGUNDA ────────────────────────────────────────────────────
    _clause_heading(doc, "CLÁUSULA SEGUNDA: OBLIGACIONES DEL CONSULTOR")
    _para_mixed(doc, [
        ("El ", False), ("CONSULTOR", True),
        (" se obliga formalmente a: a) Ejecutar los servicios contratados con estricto apego "
         "a las normativas vigentes del Instituto de Conservación Forestal (ICF) y la Ley "
         "Forestal, Áreas Protegidas y Vida Silvestre de Honduras (Decreto No. 98-2007). "
         "b) Elaborar los dictámenes, informes o planes correspondientes siguiendo las "
         "directrices éticas y técnicas de COLPROFORH. c) Entregar la documentación final "
         "en los plazos convenidos.", False),
    ])

    # ── CLÁUSULA TERCERA ────────────────────────────────────────────────────
    _clause_heading(doc, "CLÁUSULA TERCERA: OBLIGACIONES DEL PROPIETARIO")
    _para_mixed(doc, [
        ("El ", False), ("PROPIETARIO", True),
        (" se obliga a: a) Proporcionar de manera oportuna toda la documentación legal de "
         "la propiedad necesaria (título de propiedad, escrituras públicas, planos de "
         "colindancias). b) Brindar las facilidades de acceso al predio para el levantamiento "
         "de datos de campo y cubicaciones. c) Cancelar los honorarios pactados en la forma "
         "y plazos establecidos en el presente contrato.", False),
    ])

    # ── CLÁUSULA CUARTA ─────────────────────────────────────────────────────
    _clause_heading(doc, "CLÁUSULA CUARTA: VALOR DEL CONTRATO Y RETENCIONES")
    _para_mixed(doc, [
        ("Los honorarios profesionales totales acordados por la prestación del servicio "
         "forestal ascienden a la cantidad neta de ", False),
        (f"{monto_str} Lempiras", True),
        (". Este monto no incluye los gastos de tasas de inspección oficial o cobros "
         "administrativos del ICF, los cuales correrán por cuenta del ", False),
        ("PROPIETARIO", True),
        (". Ambas partes acuerdan que el ", False),
        ("PROPIETARIO", True),
        (" efectuará la retención del Impuesto sobre la Renta (ISR) conforme a la tasa "
         "del 12.5% de conformidad al Artículo 50 de la Ley del Impuesto sobre la Renta "
         "de Honduras, entregando al ", False),
        ("CONSULTOR", True),
        (" la respectiva constancia de retención en un plazo máximo de quince (15) días "
         "posteriores al pago.", False),
    ])

    # ── CLÁUSULA QUINTA ─────────────────────────────────────────────────────
    _clause_heading(doc, "CLÁUSULA QUINTA: FORMA DE PAGO")
    _para_mixed(doc, [
        ("El pago por los servicios profesionales se realizará de la siguiente manera: ", False),
        (forma_p, True), (".", False),
    ])

    # ── CLÁUSULA SEXTA ──────────────────────────────────────────────────────
    _clause_heading(doc, "CLÁUSULA SEXTA: PLAZO DE ENTREGA")
    _para_mixed(doc, [
        ("El plazo de ejecución de los servicios objeto de este contrato será de ", False),
        (plazo, True),
        (", contados a partir de la firma del presente instrumento y de la entrega completa "
         "de la información catastral e inicial por parte del ", False),
        ("PROPIETARIO", True), (".", False),
    ])

    # ── CLÁUSULA SÉPTIMA ────────────────────────────────────────────────────
    _clause_heading(doc, "CLÁUSULA SÉPTIMA: CONCILIACIÓN Y ARBITRAJE")
    _para_mixed(doc, [
        ("Cualquier diferencia o conflicto derivado de la interpretación, ejecución o "
         "rescisión de este contrato se someterá en primera instancia al procedimiento "
         "de conciliación amigable ante la junta directiva o tribunal de honor de COLPROFORH. "
         "En caso de persistir el conflicto, se someterán a la jurisdicción de los tribunales "
         "de la República de Honduras.", False),
    ])

    # ── CLÁUSULA OCTAVA (adicional, opcional) ───────────────────────────────
    cl_num_aceptacion = "OCTAVA"
    if clausulas_extra and clausulas_extra.strip():
        _clause_heading(doc,
            "CLÁUSULA OCTAVA: CONDICIONES ESPECIALES Y ACUERDOS ADICIONALES")
        _para_mixed(doc, [(clausulas_extra.strip(), False)])
        cl_num_aceptacion = "NOVENA"

    # ── CLÁUSULA ACEPTACIÓN ─────────────────────────────────────────────────
    _clause_heading(doc, f"CLÁUSULA {cl_num_aceptacion}: ACEPTACIÓN")
    _para_mixed(doc, [
        ("En fe de lo cual, y estando de acuerdo con todas y cada una de las cláusulas "
         "aquí redactadas, firmamos el presente contrato en duplicado en fecha: ", False),
        (fecha, True), (".", False),
    ], space_after=20)

    # ── Firmas ───────────────────────────────────────────────────────────────
    _add_hr(doc)

    # Tabla de firmas 2 columnas
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    # Limpiar bordes de la tabla
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right",
                         "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    def _firma_cell(cell, linea1, linea2, linea3):
        cell.paragraphs[0].clear()
        # Espacio
        p_esp = cell.add_paragraph("")
        p_esp.paragraph_format.space_after = Pt(28)
        # Línea horizontal de firma
        p_hr = cell.add_paragraph()
        pPr = p_hr._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top_e = OxmlElement("w:top")
        top_e.set(qn("w:val"), "single")
        top_e.set(qn("w:sz"), "4")
        top_e.set(qn("w:space"), "1")
        top_e.set(qn("w:color"), "999999")
        pBdr.append(top_e)
        pPr.append(pBdr)
        p_hr.paragraph_format.space_after = Pt(2)
        # Nombre/rol
        for txt, bold in [(linea1, True), (linea2, False), (linea3, False)]:
            p = cell.add_paragraph(txt)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after  = Pt(1)
            r = p.runs[0]
            r.font.size  = Pt(9)
            r.font.bold  = bold
            r.font.name  = "Times New Roman"
            r.font.color.rgb = _NEGRO

    _firma_cell(tbl.cell(0, 0),
                "EL CONSULTOR",
                c_nom,
                c_reg)
    _firma_cell(tbl.cell(0, 1),
                "EL PROPIETARIO / CLIENTE",
                cl_nom,
                f"RTN: {rtn_display}")

    doc.save(ruta_docx)
    return ruta_docx
